from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = "/Users/thunderboy/Docs/syntaxWorkspacee/nweb/AlienMint_Client_Overview.docx"
NAVY=RGBColor(7,22,27); EMERALD=RGBColor(16,185,129); CYAN=RGBColor(8,145,178); MUTED=RGBColor(83,98,104); LIGHT="E8F7F2"; INK=RGBColor(20,30,34)

def font(run, size=11, bold=False, color=INK, name="Aptos"):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),name); run._element.rPr.rFonts.set(qn("w:hAnsi"),name)
    run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=color

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); tcPr.append(shd)

def margins(cell, top=120, start=160, bottom=120, end=160):
    tc=cell._tc.get_or_add_tcPr(); tcMar=tc.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if tcMar.getparent() is None: tc.append(tcMar)
    for side,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        el=OxmlElement(f"w:{side}"); el.set(qn("w:w"),str(val)); el.set(qn("w:type"),"dxa"); tcMar.append(el)

def set_table_widths(table, widths):
    table.autofit=False
    grid=table._tbl.tblGrid
    for old in list(grid): grid.remove(old)
    for width in widths:
        col=OxmlElement("w:gridCol"); col.set(qn("w:w"),str(width)); grid.append(col)
    pr=table._tbl.tblPr; tw=pr.first_child_found_in("w:tblW"); tw.set(qn("w:w"),str(sum(widths))); tw.set(qn("w:type"),"dxa")
    ind=OxmlElement("w:tblInd"); ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa"); pr.append(ind)
    for row in table.rows:
        for cell,width in zip(row.cells,widths):
            tcw=cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW"); tcw.set(qn("w:w"),str(width)); tcw.set(qn("w:type"),"dxa")

def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f"Heading {level}"); p.add_run(text); return p

def body(doc, text, bold_lead=None):
    p=doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead): font(p.add_run(bold_lead),bold=True); font(p.add_run(text[len(bold_lead):]))
    else: font(p.add_run(text))
    return p

def bullet(doc, text):
    p=doc.add_paragraph(style="List Bullet"); font(p.add_run(text)); return p

def callout(doc, label, text, fill=LIGHT):
    t=doc.add_table(rows=1,cols=1); set_table_widths(t,[9360]); c=t.cell(0,0); shade(c,fill); margins(c,180,220,180,220)
    p=c.paragraphs[0]; font(p.add_run(label+"  "),10,bold=True,color=CYAN); font(p.add_run(text),10.5,color=INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.492)
styles=doc.styles
normal=styles["Normal"]; normal.font.name="Aptos"; normal.font.size=Pt(11); normal.font.color.rgb=INK; normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,before,after,color in [("Heading 1",16,16,8,NAVY),("Heading 2",13,12,6,CYAN),("Heading 3",12,8,4,NAVY)]:
    s=styles[name]; s.font.name="Aptos Display"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=color; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for n in ["List Bullet","List Number"]:
    s=styles[n]; s.font.name="Aptos"; s.font.size=Pt(11); s.paragraph_format.left_indent=Inches(.5); s.paragraph_format.first_line_indent=Inches(-.25); s.paragraph_format.space_after=Pt(8); s.paragraph_format.line_spacing=1.167

# Running furniture
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(hp.add_run("ALIENMINT  |  CLIENT OVERVIEW"),8.5,bold=True,color=MUTED)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(fp.add_run("Interactive product demonstration • July 2026"),8,color=MUTED)

# Cover / customer-pack opening
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(22); p.paragraph_format.space_after=Pt(3); font(p.add_run("CLIENT PRODUCT OVERVIEW"),10,bold=True,color=EMERALD)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(7); font(p.add_run("AlienMint"),31,bold=True,color=NAVY,name="Aptos Display")
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(22); font(p.add_run("A premium NFT minting experience that clients can explore without a wallet, cryptocurrency, or technical setup."),14,color=MUTED)
callout(doc,"THE SIMPLE VERSION", "AlienMint lets visitors discover an NFT collection, preview artwork, choose a quantity, and complete a realistic simulated mint. A separate live testnet version demonstrates the real blockchain workflow when credentials are configured.")

heading(doc,"What you can experience today")
body(doc,"The shared app is designed so a reviewer can open one link and immediately understand the product. No account, wallet, test ETH, or instructions are required for the main demonstration.")
for x in ["Browse a premium collection homepage and curated artwork gallery.","Choose an NFT quantity and complete a clearly labeled simulated mint.","See confirmation, processing, success, artwork reveal, and reset states.","Open Creator Studio to understand how a collection is prepared.","Open the live testnet route to review the real wallet and blockchain experience."]: bullet(doc,x)

heading(doc,"The three parts of AlienMint")
t=doc.add_table(rows=1,cols=3); set_table_widths(t,[2160,3600,3600]); t.style="Table Grid"
for c,txt in zip(t.rows[0].cells,["Area","Purpose","What the client does"]): shade(c,"DDF5EE"); margins(c); font(c.paragraphs[0].add_run(txt),10,bold=True,color=NAVY)
rows=[("Demo homepage","A polished, risk-free preview of the customer experience.","Browse artwork and complete a simulated mint."),("Creator Studio","A local preview of collection setup and metadata creation.","Select images, set collection details, and inspect metadata."),("Live testnet","The real Web3 interaction path on Base Sepolia.","Connect a wallet and mint a test NFT when configured.")]
for row in rows:
    cells=t.add_row().cells
    for c,txt in zip(cells,row): margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; font(c.paragraphs[0].add_run(txt),9.5,bold=(c==cells[0]))

doc.add_page_break()
heading(doc,"How the artwork and NFTs are added")
body(doc,"Creator Studio gives a non-technical preview of this process. During the demo, selected files stay on the user’s device and are not transmitted anywhere.")
steps=[("1. Prepare the artwork","The creator supplies one finished image for each unique NFT, or uses a generative-art pipeline to create a large collection."),("2. Add collection information","The name, symbol, description, price, supply, mint limit, external website, royalty planning details, and traits are entered."),("3. Generate NFT metadata","AlienMint creates a standard metadata document for every token. This document connects the token name, image, description, and traits."),("4. Publish to decentralized storage","After approval, artwork and metadata can be uploaded to IPFS through a service such as Pinata or an NFT.Storage-compatible provider."),("5. Deploy and verify the contract","The smart contract is deployed to the chosen blockchain, verified publicly, and connected to the mint website."),("6. Open the mint","Visitors connect their wallets, pay the configured price, and receive NFTs from the real smart contract.")]
for title,text in steps:
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(2); font(p.add_run(title),11,bold=True,color=NAVY)
    body(doc,text)
callout(doc,"IMPORTANT", "A 10,000-item unique collection requires 10,000 prepared artworks and metadata records, or a dedicated generative pipeline. The current studio intentionally uses a manageable sample batch for client approval.","FFF4D9")

heading(doc,"What is real and what is simulated")
heading(doc,"Interactive demo",2)
body(doc,"The homepage behaves like a complete mint experience, but it does not send a blockchain transaction or use real money. Every simulated action is visibly labeled as a demo.")
heading(doc,"Live testnet version",2)
body(doc,"The /live area uses the actual wallet and smart-contract flow on Base Sepolia. Testnet tokens have no real monetary value, making this suitable for technical acceptance testing before a mainnet launch.")
doc.add_page_break()
heading(doc,"Security and trust")
body(doc,"AlienMint is built with a modern Next.js interface and a Solidity ERC-721 smart contract. The contract uses ERC721A for efficient minting and includes fixed supply enforcement, exact payment checks, transaction limits, safe minting, protected withdrawals, and automated tests.")
for x in ["The demo never requests a wallet or funds.","The studio does not upload selected artwork in demo mode.","The live area clearly identifies Base Sepolia as a test network.","Contract values such as price and supply are read from the blockchain in live mode.","Production secrets and deployment keys are not placed in browser code."]: bullet(doc,x)

heading(doc,"What is needed after approval")
for x in ["Final brand name, logo, collection copy, artwork, and visual direction.","Final supply, mint price, per-transaction limit, sale timing, and supported network.","A storage provider account for permanent artwork and metadata hosting.","A secure owner wallet or multisignature wallet for contract ownership.","RPC and WalletConnect project credentials for the live website.","Contract deployment, public verification, domain setup, and end-to-end acceptance testing."]: bullet(doc,x)

heading(doc,"Client approval checklist")
for x in ["The homepage feels premium and matches the intended audience.","The simulated mint journey is clear and convincing.","The artwork gallery and collection story feel appropriate.","Creator Studio explains the upload and publishing workflow clearly.","The proposed collection settings and launch pathway are acceptable.","AlienMint may proceed to storage integration, deployment, and production hardening."]:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.2); p.paragraph_format.space_after=Pt(6); font(p.add_run("☐  "+x))

callout(doc,"RECOMMENDED NEXT STEP", "Review the demo homepage and Creator Studio first. Once the experience and collection direction are approved, move to authenticated IPFS uploads, final contract deployment, verification, and a controlled testnet acceptance session.")

# Metadata and save
doc.core_properties.title="AlienMint Client Product Overview"; doc.core_properties.subject="Plain-language overview of the AlienMint NFT application"; doc.core_properties.author="AlienMint"
doc.save(OUT)
print(OUT)
